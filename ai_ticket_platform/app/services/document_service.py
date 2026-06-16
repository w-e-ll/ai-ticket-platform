from __future__ import annotations

import shutil
from pathlib import Path
from uuid import UUID, uuid4

from docx import Document as DocxDocument
from fastapi import UploadFile
from pypdf import PdfReader
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from ai_ticket_platform.app.ai.chunking import TextChunker
from ai_ticket_platform.app.ai.embeddings import get_embedding_provider
from ai_ticket_platform.app.ai.vector_store import (
    VectorDocument,
    build_vector_document,
    get_vector_store,
)
from ai_ticket_platform.app.config import get_settings
from ai_ticket_platform.app.models.db import Document, DocumentChunk
from ai_ticket_platform.app.utils.errors import (
    DocumentParsingError,
    DocumentValidationError,
    FileStorageError,
)
from ai_ticket_platform.app.utils.logging import get_logger


logger = get_logger(__name__)


class DocumentService:
    """Service for document upload, parsing, chunking, and indexing."""

    allowed_extensions = {".pdf", ".docx", ".txt"}

    def __init__(self, db_session: AsyncSession) -> None:
        """Initialize document service."""
        self.db_session = db_session
        self.settings = get_settings()
        self.chunker = TextChunker()
        self.embedding_provider = get_embedding_provider()
        self.vector_store = get_vector_store()

        logger.info(
            "Document service initialized",
            extra={
                "event": "document_service_initialized",
                "operation": "document_service_init",
                "status": "success",
            },
        )

    async def upload_and_index_document(
        self,
        *,
        tenant_id: UUID,
        upload_file: UploadFile,
        department: str | None = None,
    ) -> Document:
        """Upload document, persist metadata, extract chunks, and index embeddings."""
        document_id = uuid4()

        logger.info(
            "Document upload and indexing started",
            extra={
                "event": "document_upload_indexing_started",
                "operation": "upload_and_index_document",
                "status": "started",
                "tenant_id": str(tenant_id),
                "document_id": str(document_id),
                "document_filename": upload_file.filename,
                "department": department,
            },
        )

        try:
            self._validate_upload_file(upload_file)

            stored_path = self._store_upload_file(
                tenant_id=tenant_id,
                document_id=document_id,
                upload_file=upload_file,
            )

            file_size_bytes = stored_path.stat().st_size
            content_type = upload_file.content_type or "application/octet-stream"

            document = Document(
                id=document_id,
                tenant_id=tenant_id,
                filename=upload_file.filename or stored_path.name,
                file_path=str(stored_path),
                content_type=content_type,
                file_size_bytes=file_size_bytes,
                department=department,
                is_processed=False,
                chunk_count=0,
                metadata_json={
                    "original_filename": upload_file.filename,
                    "extension": stored_path.suffix.lower(),
                },
            )

            self.db_session.add(document)
            await self.db_session.flush()

            logger.info(
                "Document metadata persisted",
                extra={
                    "event": "document_metadata_persisted",
                    "operation": "upload_and_index_document",
                    "status": "success",
                    "tenant_id": str(tenant_id),
                    "document_id": str(document_id),
                    "document_filename": document.filename,
                    "file_size_bytes": file_size_bytes,
                },
            )

            text = self.extract_text(stored_path)
            chunks = self.chunker.split(text)

            if not chunks:
                raise DocumentParsingError(
                    "Document does not contain extractable text",
                    details={"document_id": str(document_id)},
                )

            embeddings = self.embedding_provider.embed_texts(
                [chunk.content for chunk in chunks]
            )

            vector_documents: list[VectorDocument] = []

            for chunk, embedding in zip(chunks, embeddings, strict=True):
                db_chunk = DocumentChunk(
                    document_id=document.id,
                    chunk_index=chunk.chunk_index,
                    content=chunk.content,
                    embedding_model=self.embedding_provider.model_name(),
                    metadata_json={
                        "start_char": chunk.start_char,
                        "end_char": chunk.end_char,
                        "department": department,
                    },
                )

                self.db_session.add(db_chunk)
                await self.db_session.flush()

                vector_document = build_vector_document(
                    content=chunk.content,
                    embedding=embedding,
                    document_id=document.id,
                    chunk_id=db_chunk.id,
                    tenant_id=tenant_id,
                    department=department,
                    metadata={
                        "filename": document.filename,
                        "chunk_index": chunk.chunk_index,
                        "content_type": content_type,
                    },
                )

                db_chunk.vector_id = vector_document.id
                vector_documents.append(vector_document)

                logger.debug(
                    "Document chunk persisted",
                    extra={
                        "event": "document_chunk_persisted",
                        "operation": "upload_and_index_document",
                        "status": "success",
                        "tenant_id": str(tenant_id),
                        "document_id": str(document.id),
                        "chunk_id": str(db_chunk.id),
                        "chunk_index": chunk.chunk_index,
                    },
                )

            indexed_count = self.vector_store.add_documents(vector_documents)

            document.chunk_count = len(chunks)
            document.is_processed = True

            await self.db_session.commit()
            await self.db_session.refresh(document)

            logger.info(
                "Document upload and indexing completed",
                extra={
                    "event": "document_upload_indexing_completed",
                    "operation": "upload_and_index_document",
                    "status": "success",
                    "tenant_id": str(tenant_id),
                    "document_id": str(document.id),
                    "document_filename": document.filename,
                    "chunk_count": len(chunks),
                    "indexed_count": indexed_count,
                    "embedding_model": self.embedding_provider.model_name(),
                },
            )

            return document

        except (
            DocumentValidationError,
            DocumentParsingError,
            FileStorageError,
        ):
            await self.db_session.rollback()
            raise

        except Exception as exc:
            await self.db_session.rollback()

            logger.error(
                "Document upload and indexing failed",
                extra={
                    "event": "document_upload_indexing_failed",
                    "operation": "upload_and_index_document",
                    "status": "failed",
                    "tenant_id": str(tenant_id),
                    "document_id": str(document_id),
                    "document_filename": upload_file.filename,
                    "exception_type": type(exc).__name__,
                },
                exc_info=True,
            )

            raise DocumentParsingError(
                "Failed to upload and index document",
                details={"document_id": str(document_id)},
            ) from exc

    async def get_document(
        self,
        *,
        tenant_id: UUID,
        document_id: UUID,
    ) -> Document:
        """Return document by tenant and document id."""
        logger.info(
            "Document retrieval started",
            extra={
                "event": "document_retrieval_started",
                "operation": "get_document",
                "status": "started",
                "tenant_id": str(tenant_id),
                "document_id": str(document_id),
            },
        )

        result = await self.db_session.execute(
            select(Document).where(
                Document.id == document_id,
                Document.tenant_id == tenant_id,
            )
        )

        document = result.scalar_one_or_none()

        if document is None:
            logger.warning(
                "Document not found",
                extra={
                    "event": "document_not_found",
                    "operation": "get_document",
                    "status": "warning",
                    "tenant_id": str(tenant_id),
                    "document_id": str(document_id),
                },
            )

            raise DocumentValidationError(
                "Document not found",
                details={
                    "tenant_id": str(tenant_id),
                    "document_id": str(document_id),
                },
            )

        logger.info(
            "Document retrieval completed",
            extra={
                "event": "document_retrieval_completed",
                "operation": "get_document",
                "status": "success",
                "tenant_id": str(tenant_id),
                "document_id": str(document_id),
            },
        )

        return document

    async def list_documents(
        self,
        *,
        tenant_id: UUID,
        limit: int = 50,
        offset: int = 0,
    ) -> list[Document]:
        """List tenant documents."""
        if limit < 1 or limit > 200:
            raise DocumentValidationError("limit must be between 1 and 200")

        if offset < 0:
            raise DocumentValidationError("offset must not be negative")

        logger.info(
            "Document listing started",
            extra={
                "event": "document_listing_started",
                "operation": "list_documents",
                "status": "started",
                "tenant_id": str(tenant_id),
                "limit": limit,
                "offset": offset,
            },
        )

        result = await self.db_session.execute(
            select(Document)
            .where(Document.tenant_id == tenant_id)
            .order_by(Document.created_at.desc())
            .limit(limit)
            .offset(offset)
        )

        documents = list(result.scalars().all())

        logger.info(
            "Document listing completed",
            extra={
                "event": "document_listing_completed",
                "operation": "list_documents",
                "status": "success",
                "tenant_id": str(tenant_id),
                "document_count": len(documents),
            },
        )

        return documents

    def extract_text(self, file_path: Path) -> str:
        """Extract plain text from supported document type."""
        logger.info(
            "Document text extraction started",
            extra={
                "event": "document_text_extraction_started",
                "operation": "extract_text",
                "status": "started",
                "document_filename": file_path.name,
            },
        )

        try:
            extension = file_path.suffix.lower()

            if extension == ".pdf":
                text = self._extract_pdf_text(file_path)
            elif extension == ".docx":
                text = self._extract_docx_text(file_path)
            elif extension == ".txt":
                text = self._extract_txt_text(file_path)
            else:
                raise DocumentValidationError(
                    "Unsupported document extension",
                    details={"extension": extension},
                )

            if not text.strip():
                raise DocumentParsingError(
                    "No text could be extracted from document",
                    details={"filename": file_path.name},
                )

            logger.info(
                "Document text extraction completed",
                extra={
                    "event": "document_text_extraction_completed",
                    "operation": "extract_text",
                    "status": "success",
                    "document_filename": file_path.name,
                    "text_length": len(text),
                },
            )

            return text

        except (DocumentValidationError, DocumentParsingError):
            raise

        except Exception as exc:
            logger.error(
                "Document text extraction failed",
                extra={
                    "event": "document_text_extraction_failed",
                    "operation": "extract_text",
                    "status": "failed",
                    "document_filename": file_path.name,
                    "exception_type": type(exc).__name__,
                },
                exc_info=True,
            )

            raise DocumentParsingError(
                "Failed to extract document text",
                details={"filename": file_path.name},
            ) from exc

    def _validate_upload_file(self, upload_file: UploadFile) -> None:
        """Validate uploaded file metadata."""
        filename = upload_file.filename

        if not filename:
            raise DocumentValidationError("uploaded file must have a filename")

        extension = Path(filename).suffix.lower()

        if extension not in self.allowed_extensions:
            raise DocumentValidationError(
                "Unsupported file extension",
                details={
                    "filename": filename,
                    "extension": extension,
                    "allowed_extensions": sorted(self.allowed_extensions),
                },
            )

        logger.debug(
            "Upload file validation completed",
            extra={
                "event": "upload_file_validation_completed",
                "operation": "validate_upload_file",
                "status": "success",
                "document_filename": filename,
                "extension": extension,
            },
        )

    def _store_upload_file(
        self,
        *,
        tenant_id: UUID,
        document_id: UUID,
        upload_file: UploadFile,
    ) -> Path:
        """Persist uploaded file to local storage."""
        try:
            filename = upload_file.filename or f"{document_id}.bin"
            extension = Path(filename).suffix.lower()

            tenant_upload_dir = (
                self.settings.upload_dir
                / str(tenant_id)
                / str(document_id)
            )
            tenant_upload_dir.mkdir(parents=True, exist_ok=True)

            stored_path = tenant_upload_dir / f"source{extension}"

            logger.info(
                "Upload file storage started",
                extra={
                    "event": "upload_file_storage_started",
                    "operation": "store_upload_file",
                    "status": "started",
                    "tenant_id": str(tenant_id),
                    "document_id": str(document_id),
                    "document_filename": filename,
                },
            )

            with stored_path.open("wb") as output_file:
                upload_file.file.seek(0)
                shutil.copyfileobj(upload_file.file, output_file)

            file_size_bytes = stored_path.stat().st_size
            max_size_bytes = self.settings.max_file_size_mb * 1024 * 1024

            if file_size_bytes <= 0:
                raise FileStorageError("uploaded file is empty")

            if file_size_bytes > max_size_bytes:
                stored_path.unlink(missing_ok=True)
                raise FileStorageError(
                    "uploaded file exceeds maximum size",
                    details={
                        "file_size_bytes": file_size_bytes,
                        "max_size_bytes": max_size_bytes,
                    },
                )

            logger.info(
                "Upload file storage completed",
                extra={
                    "event": "upload_file_storage_completed",
                    "operation": "store_upload_file",
                    "status": "success",
                    "tenant_id": str(tenant_id),
                    "document_id": str(document_id),
                    "document_filename": filename,
                    "file_size_bytes": file_size_bytes,
                },
            )

            return stored_path

        except FileStorageError:
            raise

        except Exception as exc:
            logger.error(
                "Upload file storage failed",
                extra={
                    "event": "upload_file_storage_failed",
                    "operation": "store_upload_file",
                    "status": "failed",
                    "tenant_id": str(tenant_id),
                    "document_id": str(document_id),
                    "document_filename": upload_file.filename,
                    "exception_type": type(exc).__name__,
                },
                exc_info=True,
            )

            raise FileStorageError("Failed to store uploaded file") from exc

    def _extract_pdf_text(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        reader = PdfReader(str(file_path))
        page_texts: list[str] = []

        for page_index, page in enumerate(reader.pages):
            text = page.extract_text() or ""

            logger.debug(
                "PDF page text extracted",
                extra={
                    "event": "pdf_page_text_extracted",
                    "operation": "extract_pdf_text",
                    "status": "success",
                    "document_filename": file_path.name,
                    "page_index": page_index,
                    "text_length": len(text),
                },
            )

            page_texts.append(text)

        return "\n".join(page_texts)

    def _extract_docx_text(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        document = DocxDocument(str(file_path))

        paragraphs = [
            paragraph.text.strip()
            for paragraph in document.paragraphs
            if paragraph.text and paragraph.text.strip()
        ]

        logger.debug(
            "DOCX text extracted",
            extra={
                "event": "docx_text_extracted",
                "operation": "extract_docx_text",
                "status": "success",
                "document_filename": file_path.name,
                "paragraph_count": len(paragraphs),
            },
        )

        return "\n".join(paragraphs)

    def _extract_txt_text(self, file_path: Path) -> str:
        """Extract text from TXT file."""
        text = file_path.read_text(encoding="utf-8")

        logger.debug(
            "TXT text extracted",
            extra={
                "event": "txt_text_extracted",
                "operation": "extract_txt_text",
                "status": "success",
                "document_filename": file_path.name,
                "text_length": len(text),
            },
        )

        return text
